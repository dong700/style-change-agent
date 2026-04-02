"""
文章风格分析 Web 应用 - Flask 后端
"""

import os
import json
import sys
from flask import Flask, request, jsonify, send_from_directory, render_template, send_file
from werkzeug.utils import secure_filename
from datetime import datetime

# 添加项目根目录到路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from style_extractor import analyze_document, extract_text_from_docx
from style_rewriter import StyleRewriter

app = Flask(__name__, 
            template_folder='../templates',
            static_folder='../static')

# 配置
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 最大 16MB
app.config['ALLOWED_EXTENSIONS'] = {'docx'}

# 确保目录存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

def allowed_file(filename):
    """检查文件扩展名是否允许"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/')
def index():
    """首页"""
    return render_template('index.html')

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传并分析文件"""
    # 检查是否有文件
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': '没有选择文件'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': '只支持 .docx 格式的文件'}), 400
    
    try:
        # 保存上传的文件
        filename = secure_filename(file.filename)
        # 添加时间戳避免重名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # 分析文档
        output_dir = os.path.join(app.config['OUTPUT_FOLDER'], os.path.splitext(filename)[0])
        os.makedirs(output_dir, exist_ok=True)
        
        # 分析文档（使用 LLM）
        features = analyze_document(filepath, use_llm=True, output_dir=output_dir)
        
        # 读取 JSON 结果
        json_path = features.get('json_path', '')
        if os.path.exists(json_path):
            with open(json_path, 'r', encoding='utf-8') as f:
                result = json.load(f)
        else:
            result = features
        
        # 添加文件信息
        result['upload_filename'] = filename
        result['original_filename'] = file.filename
        
        return jsonify({
            'success': True,
            'data': result,
            'message': '分析成功'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/modify_style', methods=['POST'])
def modify_style():
    """修改风格参数"""
    try:
        data = request.json
        
        # 获取修改后的风格参数
        modified_features = data.get('features', {})
        
        # 这里可以添加逻辑来根据修改后的风格参数生成新的文本
        # 目前是返回修改后的参数
        return jsonify({
            'success': True,
            'data': modified_features,
            'message': '风格参数已更新'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/documents')
def list_documents():
    """列出所有已分析的文档"""
    try:
        documents = []
        output_folder = app.config['OUTPUT_FOLDER']
        
        if os.path.exists(output_folder):
            for folder in os.listdir(output_folder):
                folder_path = os.path.join(output_folder, folder)
                if os.path.isdir(folder_path):
                    # 查找 JSON 文件
                    for file in os.listdir(folder_path):
                        if file.endswith('.json'):
                            json_path = os.path.join(folder_path, file)
                            with open(json_path, 'r', encoding='utf-8') as f:
                                doc_data = json.load(f)
                                documents.append({
                                    'id': folder,
                                    'filename': doc_data.get('基本信息', {}).get('文件名', file),
                                    'analyzed_at': folder.split('_')[0] if '_' in folder else 'Unknown'
                                })
                            break
        
        return jsonify({
            'success': True,
            'documents': documents
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/documents/<doc_id>')
def get_document(doc_id):
    """获取单个文档的详细信息"""
    try:
        output_folder = app.config['OUTPUT_FOLDER']
        doc_path = os.path.join(output_folder, doc_id)
        
        if not os.path.exists(doc_path):
            return jsonify({'error': '文档不存在'}), 404
        
        # 查找 JSON 文件
        for file in os.listdir(doc_path):
            if file.endswith('.json'):
                json_path = os.path.join(doc_path, file)
                with open(json_path, 'r', encoding='utf-8') as f:
                    doc_data = json.load(f)
                
                return jsonify({
                    'success': True,
                    'data': doc_data
                })
        
        return jsonify({'error': '未找到分析结果'}), 404
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/rewrite', methods=['POST'])
def rewrite_text():
    """改写文本"""
    try:
        data = request.json
        
        # 获取参数
        target_text = data.get('target_text', '')
        style_labels = data.get('style_labels', [])
        style_words = data.get('style_words', {})
        quantitative_features = data.get('quantitative_features', {})
        
        if not target_text:
            return jsonify({'error': '请提供目标文本'}), 400
        
        if not style_labels:
            return jsonify({'error': '请提供风格标签'}), 400
        
        # 创建改写器
        rewriter = StyleRewriter(model='qwen-max')
        
        # 执行改写（传入量化特征）
        result = rewriter.rewrite_with_comparison(
            target_text, 
            style_labels, 
            style_words,
            quantitative_features
        )
        
        return jsonify({
            'success': True,
            'data': result,
            'message': '改写成功'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/rewrite_from_file', methods=['POST'])
def rewrite_from_file():
    """从上传的 Word 文档改写"""
    try:
        # 检查是否有文件
        if 'file' not in request.files:
            return jsonify({'error': '没有上传文件'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': '没有选择文件'}), 400
        
        # 获取风格参数
        style_labels_str = request.form.get('style_labels')
        style_words_str = request.form.get('style_words')
        quantitative_features_str = request.form.get('quantitative_features')
        
        if not style_labels_str:
            return jsonify({'error': '请提供风格标签'}), 400
        
        # 解析 JSON 参数
        try:
            style_labels = json.loads(style_labels_str)
            style_words = json.loads(style_words_str) if style_words_str else {}
            quantitative_features = json.loads(quantitative_features_str) if quantitative_features_str else {}
        except json.JSONDecodeError as e:
            return jsonify({'error': f'参数解析失败：{str(e)}'}), 400
        
        # 保存上传的文件
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # 提取文本
        target_text = extract_text_from_docx(filepath)
        
        if not target_text:
            return jsonify({'error': '无法从文档中提取文本'}), 400
        
        # 创建改写器
        rewriter = StyleRewriter(model='qwen-max')
        
        # 执行改写
        result = rewriter.rewrite_with_comparison(
            target_text,
            style_labels,
            style_words,
            quantitative_features
        )
        
        # 添加文件信息
        result['upload_filename'] = filename
        result['original_filename'] = file.filename
        
        return jsonify({
            'success': True,
            'data': result,
            'message': '改写成功'
        })
        
    except Exception as e:
        print(f"改写失败：{e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/export_word', methods=['POST'])
def export_word():
    """导出改写结果为 Word 文档"""
    try:
        data = request.json
        
        # 获取参数
        rewritten_text = data.get('rewritten_text', '')
        original_text = data.get('original_text', '')
        original_filename = data.get('filename', '改写结果')
        
        if not rewritten_text:
            return jsonify({'error': '没有可导出的内容'}), 400
        
        # 创建 Word 文档
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # 设置默认字体（中文）
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加标题
        heading = doc.add_heading('文章改写结果', level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加基本信息
        doc.add_paragraph()
        info_para = doc.add_paragraph()
        info_run = info_para.add_run(f'原文文件名：{original_filename}\n')
        info_run.bold = True
        info_run.font.name = '宋体'
        info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        time_run = info_para.add_run(f'导出时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        time_run.italic = True
        time_run.font.name = '宋体'
        time_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加改写后的内容
        doc.add_page_break()
        doc.add_heading('改写后内容', level=1)
        
        # 分段添加改写后的文本，保持段落结构
        paragraphs = rewritten_text.split('\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph()
                run = p.add_run(para)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 保存文件
        output_filename = f"{os.path.splitext(original_filename)[0]}_改写版.docx"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        doc.save(output_path)
        
        # 发送文件
        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"导出失败：{e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'导出失败：{str(e)}'
        }), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
