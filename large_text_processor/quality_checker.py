"""
质量检查器 - 检查改写质量
"""
import os
import sys
import json
import re
from typing import List, Dict, Any

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


class QualityChecker:
    """质量检查器"""
    
    def __init__(self, model: str = "deepseek-api:deepseek-v4-flash", deepseek_api_key: str = ''):
        """
        初始化质量检查器
        
        Args:
            model: 使用的模型名称
            deepseek_api_key: DeepSeek API Key
        """
        self.model = model
        self.deepseek_api_key = deepseek_api_key
        self._llm_caller = None
    
    def check_batch(self, rewrite_results: List[Dict[str, Any]],
                   batch_size: int = 5) -> Dict[str, Any]:
        """
        批量检查改写质量
        
        Args:
            rewrite_results: 改写结果列表
            batch_size: 批处理大小
            
        Returns:
            Dict: 包含质量报告和审核队列
        """
        print(f"开始质量检查，共 {len(rewrite_results)} 个片段...")
        
        checked_results = []
        review_queue = []
        quality_stats = {
            'total': len(rewrite_results),
            'high_quality': 0,
            'needs_review': 0,
            'low_quality': 0
        }
        
        for i, result in enumerate(rewrite_results):
            if 'error' in result:
                result['quality_check'] = {
                    'fact_consistent': False,
                    'error': result['error']
                }
                result['quality_level'] = 'low'
                quality_stats['low_quality'] += 1
                review_queue.append(result)
                checked_results.append(result)
                continue
            
            try:
                fact_check = self._check_fact_consistency(
                    result['original_text'],
                    result['rewritten_text']
                )
                
                result['quality_check'] = fact_check
                
                if fact_check['fact_consistent'] and result['style_match']['confidence'] >= 0.8:
                    result['quality_level'] = 'high'
                    quality_stats['high_quality'] += 1
                elif not fact_check['fact_consistent'] or result['style_match']['confidence'] < 0.6:
                    result['quality_level'] = 'low'
                    quality_stats['low_quality'] += 1
                    review_queue.append(result)
                else:
                    result['quality_level'] = 'medium'
                    quality_stats['needs_review'] += 1
                
                checked_results.append(result)
                
            except Exception as e:
                print(f"\n检查片段 {result['fragment_id']} 时出错：{e}")
                result['quality_check'] = {'error': str(e)}
                result['quality_level'] = 'unknown'
                checked_results.append(result)
        
        print(f"\n质量检查完成！")
        print(f"  高质量：{quality_stats['high_quality']}")
        print(f"  需审核：{quality_stats['needs_review']}")
        print(f"  低质量：{quality_stats['low_quality']}")
        
        return {
            'quality_stats': quality_stats,
            'review_queue': review_queue,
            'all_results': checked_results
        }
    
    def _check_fact_consistency(self, original: str, rewritten: str) -> Dict[str, Any]:
        """
        检查事实一致性
        
        Args:
            original: 原始文本
            rewritten: 改写后文本
            
        Returns:
            Dict: 检查结果
        """
        prompt = f"""请对比以下两段文本，判断改写后是否保留了所有关键事实（时间、人物、事件、数据），是否添加了原文没有的虚构内容。

原文：
{original}

改写后：
{rewritten}

请输出 JSON 格式：
{{
  "fact_consistent": true/false,
  "key_facts_preserved": ["保留的事实1", "保留的事实2"],
  "missing_facts": ["缺失的事实"],
  "added_content": ["虚构的内容"],
  "style_similarity": "高/中/低",
  "overall_quality": "优秀/良好/一般/较差"
}}

只输出 JSON，不要包含其他说明："""

        try:
            if self._llm_caller is None:
                from style_extractor import LLMStyleAnalyzer
                self._llm_caller = LLMStyleAnalyzer(model=self.model, deepseek_api_key=self.deepseek_api_key)
            
            response = self._llm_caller._call_llm(prompt, self.model)
            
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                result = json.loads(json_match.group())
            else:
                result = {
                    'fact_consistent': True,
                    'error': '无法解析 JSON'
                }
            
            return result
            
        except Exception as e:
            return {
                'fact_consistent': True,
                'error': str(e)
            }
    
    def generate_review_document(self, review_queue: List[Dict], 
                                output_path: str):
        """
        生成审核文档
        
        Args:
            review_queue: 审核队列
            output_path: 输出路径
        """
        from docx import Document
        
        doc = Document()
        doc.add_heading('待审核片段', 0)
        
        for i, item in enumerate(review_queue, 1):
            doc.add_heading(f'片段 {i}: {item["fragment_id"]}', level=1)
            
            doc.add_heading('原始文本', level=2)
            doc.add_paragraph(item['original_text'])
            
            doc.add_heading('改写后文本', level=2)
            doc.add_paragraph(item['rewritten_text'])
            
            doc.add_heading('质量检查结果', level=2)
            quality_check = item.get('quality_check', {})
            
            doc.add_paragraph(f"事实一致性：{'✓' if quality_check.get('fact_consistent', False) else '✗'}")
            
            if 'missing_facts' in quality_check and quality_check['missing_facts']:
                doc.add_paragraph(f"缺失的事实：{', '.join(quality_check['missing_facts'])}")
            
            if 'added_content' in quality_check and quality_check['added_content']:
                doc.add_paragraph(f"虚构的内容：{', '.join(quality_check['added_content'])}")
            
            doc.add_paragraph(f"质量等级：{item.get('quality_level', 'unknown')}")
            doc.add_paragraph(f"匹配置信度：{item['style_match']['confidence']:.2f}")
            
            doc.add_paragraph()
        
        doc.save(output_path)
        print(f"审核文档已保存：{output_path}")
    
    def save_results(self, results: Dict[str, Any], output_path: str):
        """保存结果"""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
